"""
SharePoint integration via Microsoft Graph API.
Reads historical RFP documents from SharePoint for corpus ingestion
and writes generated drafts back for human review.
Uses DefaultAzureCredential — requires the managed identity to have
Sites.ReadWrite.All on the target SharePoint site.
"""
from __future__ import annotations

import io
import logging
import re
from pathlib import Path
from typing import Optional

import httpx
from azure.identity import DefaultAzureCredential
from azure.identity import get_bearer_token_provider

log = logging.getLogger(__name__)

GRAPH_BASE = "https://graph.microsoft.com/v1.0"
GRAPH_SCOPE = "https://graph.microsoft.com/.default"


class SharePointClient:
    def __init__(self, site_id: str, credential: Optional[DefaultAzureCredential] = None):
        """
        Args:
            site_id: SharePoint site ID in format {tenant}.sharepoint.com,{site-guid},{web-guid}
                     or the hostname/path form: sites/{hostname}:/sites/{site-name}
        """
        self._site_id = site_id
        self._credential = credential or DefaultAzureCredential()
        self._token_provider = get_bearer_token_provider(self._credential, GRAPH_SCOPE)

    def _headers(self) -> dict[str, str]:
        token = self._token_provider()
        return {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}

    # ── Read operations ──────────────────────────────────────────────────────

    async def list_library_files(self, library_name: str = "Documents") -> list[dict]:
        """List all files in a SharePoint document library."""
        url = f"{GRAPH_BASE}/sites/{self._site_id}/drives"
        async with httpx.AsyncClient(timeout=30) as client:
            drives_resp = await client.get(url, headers=self._headers())
            drives_resp.raise_for_status()
            drives = drives_resp.json()["value"]

        drive = next((d for d in drives if d["name"] == library_name), None)
        if not drive:
            raise ValueError(f"Library '{library_name}' not found on site {self._site_id}")

        drive_id = drive["id"]
        items_url = f"{GRAPH_BASE}/drives/{drive_id}/root/children"
        async with httpx.AsyncClient(timeout=30) as client:
            resp = await client.get(items_url, headers=self._headers())
            resp.raise_for_status()

        files = [
            {"id": item["id"], "name": item["name"], "size": item.get("size", 0),
             "web_url": item.get("webUrl"), "drive_id": drive_id,
             "mime_type": item.get("file", {}).get("mimeType")}
            for item in resp.json()["value"]
            if "file" in item
        ]
        log.info("Found %d files in SharePoint library '%s'", len(files), library_name)
        return files

    async def download_file(self, drive_id: str, item_id: str) -> bytes:
        """Download a file from SharePoint by drive and item ID."""
        url = f"{GRAPH_BASE}/drives/{drive_id}/items/{item_id}/content"
        async with httpx.AsyncClient(timeout=120, follow_redirects=True) as client:
            resp = await client.get(url, headers=self._headers())
            resp.raise_for_status()
        return resp.content

    async def download_library_to_local(
        self, library_name: str, dest_dir: Path, extensions: tuple[str, ...] = (".pdf", ".docx", ".md")
    ) -> list[Path]:
        """Download all matching files from a SharePoint library to a local directory."""
        dest_dir.mkdir(parents=True, exist_ok=True)
        files = await self.list_library_files(library_name)
        downloaded: list[Path] = []

        for f in files:
            if not any(f["name"].lower().endswith(ext) for ext in extensions):
                continue
            content = await self.download_file(f["drive_id"], f["id"])
            dest = dest_dir / f["name"]
            dest.write_bytes(content)
            downloaded.append(dest)
            log.info("Downloaded %s (%d bytes)", f["name"], len(content))

        return downloaded

    # ── Write operations ─────────────────────────────────────────────────────

    async def upload_draft(
        self, library_name: str, file_name: str, content: str, folder: str = "GeneratedDrafts"
    ) -> str:
        """Upload a generated RFP draft (markdown) to SharePoint. Returns the web URL."""
        url = f"{GRAPH_BASE}/sites/{self._site_id}/drives"
        async with httpx.AsyncClient(timeout=30) as client:
            drives_resp = await client.get(url, headers=self._headers())
            drives_resp.raise_for_status()
            drives = drives_resp.json()["value"]

        drive = next((d for d in drives if d["name"] == library_name), drives[0])
        drive_id = drive["id"]

        upload_url = f"{GRAPH_BASE}/drives/{drive_id}/root:/{folder}/{file_name}:/content"
        headers = {**self._headers(), "Content-Type": "text/markdown"}

        async with httpx.AsyncClient(timeout=60) as client:
            resp = await client.put(upload_url, headers=headers, content=content.encode())
            resp.raise_for_status()

        web_url = resp.json().get("webUrl", "")
        log.info("Draft uploaded to SharePoint: %s", web_url)
        return web_url

    async def upload_draft_docx(
        self,
        library_name: str,
        rfp_id: str,
        draft_id: str,
        sections: dict[str, str],
        folder: str = "GeneratedDrafts",
    ) -> str:
        """Assemble a Word document from RFP sections and upload to SharePoint.

        Uses the site default drive to avoid drive-listing permission issues.
        The folder is created automatically by Graph if it does not exist.
        """
        docx_bytes = SharePointClient.draft_to_docx(rfp_id, sections)
        safe_rfp_id = rfp_id.replace("/", "-").replace("\\", "-")
        file_name = f"{safe_rfp_id}_{draft_id}.docx"

        upload_url = (
            f"{GRAPH_BASE}/sites/{self._site_id}/drive"
            f"/root:/{folder}/{file_name}:/content"
        )
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        headers = {**self._headers(), "Content-Type": content_type}

        async with httpx.AsyncClient(timeout=60) as client:
            resp = await client.put(upload_url, headers=headers, content=docx_bytes)
            resp.raise_for_status()

        web_url = resp.json().get("webUrl", "")
        log.info("Word draft uploaded to SharePoint: %s", web_url)
        return web_url

    async def create_review_item(self, site_list_name: str, draft_metadata: dict) -> str:
        """Create a list item in a SharePoint review tracking list."""
        url = f"{GRAPH_BASE}/sites/{self._site_id}/lists/{site_list_name}/items"
        payload = {"fields": draft_metadata}
        async with httpx.AsyncClient(timeout=30) as client:
            resp = await client.post(url, headers=self._headers(), json=payload)
            resp.raise_for_status()
        return resp.json()["id"]

    async def ensure_review_list(self, site_list_name: str = "RFP Review Tracking") -> None:
        """Create the review-tracking list if it doesn't exist yet on the site.

        Uses the managed identity's app-only Sites.ReadWrite.All — this can create
        lists tenant-wide regardless of the deploying user's own SharePoint
        permission level on the site (list creation needs Design/Full Control,
        which a delegated install-time login may not have).
        """
        async with httpx.AsyncClient(timeout=30) as client:
            check = await client.get(
                f"{GRAPH_BASE}/sites/{self._site_id}/lists/{site_list_name}",
                headers=self._headers(),
            )
            if check.status_code == 200:
                return

            payload = {
                "displayName": site_list_name,
                "list": {"template": "genericList"},
                "columns": [
                    {"name": "DraftId", "text": {}},
                    {"name": "GateDecision", "text": {}},
                    {"name": "CompletenessScore", "number": {"decimalPlaces": "automatic"}},
                    {"name": "GroundednessScore", "number": {"decimalPlaces": "automatic"}},
                    {"name": "SharePointUrl", "text": {}},
                    {"name": "Status", "text": {}},
                ],
            }
            resp = await client.post(
                f"{GRAPH_BASE}/sites/{self._site_id}/lists",
                headers=self._headers(),
                json=payload,
            )
            resp.raise_for_status()
            log.info("Created SharePoint review-tracking list: %s", site_list_name)

    # ── Utility ───────────────────────────────────────────────────────────────

    @staticmethod
    def _add_inline_md(paragraph, text: str) -> None:
        """Add inline markdown (bold/italic) as properly formatted Word runs."""
        import re
        parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                paragraph.add_run(part[2:-2]).bold = True
            elif part.startswith("*") and part.endswith("*"):
                paragraph.add_run(part[1:-1]).italic = True
            else:
                paragraph.add_run(part)

    @staticmethod
    def _add_md_block(doc, line: str) -> None:
        """Convert a single markdown line to the appropriate Word paragraph type."""
        import re
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        line = line.rstrip()
        if not line:
            return

        heading_m = re.match(r'^(#{1,4})\s+(.+)$', line)
        if heading_m:
            level = min(len(heading_m.group(1)) + 1, 4)
            doc.add_heading(heading_m.group(2), level=level)
            return

        bullet_m = re.match(r'^[-*]\s+(.+)$', line)
        if bullet_m:
            p = doc.add_paragraph(style='List Bullet')
            SharePointClient._add_inline_md(p, bullet_m.group(1))
            return

        num_m = re.match(r'^\d+\.\s+(.+)$', line)
        if num_m:
            p = doc.add_paragraph(style='List Number')
            SharePointClient._add_inline_md(p, num_m.group(1))
            return

        # Horizontal rule — skip
        if re.match(r'^-{3,}$', line):
            return

        p = doc.add_paragraph()
        SharePointClient._add_inline_md(p, line)

    @staticmethod
    def _render_md_table(doc, table_lines: list[str]) -> None:
        """Convert markdown table lines into a Word table."""
        from docx.shared import RGBColor
        rows = [
            [cell.strip() for cell in line.strip().strip("|").split("|")]
            for line in table_lines
            if not re.match(r'^\|?[\s:|-]+\|?$', line.strip())  # skip separator row
        ]
        if not rows:
            return
        col_count = max(len(r) for r in rows)
        tbl = doc.add_table(rows=len(rows), cols=col_count)
        tbl.style = "Table Grid"
        for r_idx, row_data in enumerate(rows):
            for c_idx in range(col_count):
                cell = tbl.rows[r_idx].cells[c_idx]
                text = row_data[c_idx] if c_idx < len(row_data) else ""
                p = cell.paragraphs[0]
                SharePointClient._add_inline_md(p, text)
                if r_idx == 0:
                    for run in p.runs:
                        run.bold = True

    @staticmethod
    def _render_section(doc, content: str) -> None:
        """Render markdown content into Word paragraphs, handling tables as blocks."""
        lines = content.splitlines()
        i = 0
        while i < len(lines):
            line = lines[i]
            # Detect start of a markdown table
            if re.match(r'^\s*\|', line):
                table_block = []
                while i < len(lines) and re.match(r'^\s*\|', lines[i]):
                    table_block.append(lines[i])
                    i += 1
                SharePointClient._render_md_table(doc, table_block)
            else:
                SharePointClient._add_md_block(doc, line)
                i += 1

    @staticmethod
    def draft_to_docx(rfp_id: str, sections: dict[str, str]) -> bytes:
        """Assemble an RFP draft as a Word (.docx) document. Returns raw bytes."""
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        SECTION_HEADINGS = {
            "background":             "1. Background and Purpose",
            "funding_parameters":     "2. Funding Parameters",
            "eligibility":            "3. Eligibility Criteria",
            "scope_of_work":          "4. Scope of Work",
            "reporting_requirements": "5. Reporting Requirements",
            "budget_requirements":    "6. Budget Requirements",
            "evaluation_criteria":    "7. Evaluation Criteria",
            "submission_instructions":"8. Submission Instructions",
        }

        doc = Document()

        # Cover
        title = doc.add_heading("REQUEST FOR PROPOSALS", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sub = doc.add_paragraph(f"RFP ID: {rfp_id}")
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].italic = True

        notice = doc.add_paragraph("AI-generated draft — pending human review and approval before release.")
        notice.runs[0].font.size = Pt(9)
        notice.runs[0].italic = True

        doc.add_paragraph("")  # spacer

        for key, heading in SECTION_HEADINGS.items():
            doc.add_heading(heading, level=1)
            content = sections.get(key, "Not generated.")
            SharePointClient._render_section(doc, content)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()

    @staticmethod
    def draft_to_markdown(rfp_id: str, sections: dict[str, str]) -> str:
        """Render an RFP draft dictionary into a Markdown document."""
        header_map = {
            "background": "1. Background and Purpose",
            "funding_parameters": "2. Funding Parameters",
            "eligibility": "3. Eligibility Criteria",
            "scope_of_work": "4. Scope of Work",
            "reporting_requirements": "5. Reporting Requirements",
            "budget_requirements": "6. Budget Requirements",
            "evaluation_criteria": "7. Evaluation Criteria",
            "submission_instructions": "8. Submission Instructions",
        }
        lines = [f"# REQUEST FOR PROPOSALS\n**RFP ID:** {rfp_id} *(AI-generated draft — pending human review)*\n\n---\n"]
        for key, heading in header_map.items():
            lines.append(f"## {heading}\n\n{sections.get(key, '_Not generated_')}\n\n---\n")
        return "\n".join(lines)
